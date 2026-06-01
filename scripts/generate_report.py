"""
Generate final report (DOCX) for Assignment 3.
Run AFTER all pipeline steps have been executed.
"""

import json
import os
import sys
from pathlib import Path
from datetime import datetime

BASE_DIR = Path(__file__).resolve().parents[1]
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def load_metrics():
    path = BASE_DIR / "output" / "metrics_summary.json"
    if path.exists():
        with open(path) as f:
            return json.load(f)
    print(f"[WARN] metrics_summary.json not found at {path}")
    return {}

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)

def generate_report():
    metrics = load_metrics()
    doc = Document()

    # Title
    title = doc.add_heading("Job Market Analytics Pipeline", level=0)
    doc.add_paragraph("Assignment 3 - Data Engineering Pipeline\nTools & Techniques for Data Science\n")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    # 1. Introduction
    add_heading(doc, "1. Introduction and Problem Statement")
    add_para(doc,
        "This project implements an automated job market analytics pipeline that extracts, "
        "cleans, standardizes, filters, and analyzes AI/ML/Data-related job postings from "
        "four free online sources. The goal is to convert messy, multi-source job data into "
        "a clean, validated, analytics-ready dataset and provide actionable business insights "
        "including job counts by source, remote/on-site ratios, experience bracket distribution, "
        "salary analysis, and skill demand."
    )

    # 2. Source Selection
    add_heading(doc, "2. Source Selection and API Links")
    sources = [
        ("Arbeitnow", "https://www.arbeitnow.com/api/job-board-api", "Jobs with Europe/Germany focus; includes remote flag"),
        ("RemoteOK", "https://remoteok.com/api", "Remote-only job listings; includes salary data"),
        ("Himalayas", "https://himalayas.app/jobs/api/search?q=data&sort=recent&page=1", "Remote jobs with filters for keyword, seniority, employment type"),
        ("RemoteJobs.org", "https://remotejobs.org/api/v1/jobs?category=data-science&limit=50", "Categorized job listings with salary fields"),
    ]
    for name, url, note in sources:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(f"{url} - {note}")

    # 3. Pipeline Architecture
    add_heading(doc, "3. Pipeline Architecture")
    add_para(doc,
        "The pipeline follows a layered architecture: "
    )
    steps = [
        "Extraction Layer: 4 Python scripts fetch data from each source in parallel",
        "Merging Layer: Schema standardization and source unification",
        "Cleaning Layer (KNIME): HTML removal, deduplication, AI/ML filtering, skill extraction, salary and experience processing",
        "Validation Layer: Data quality checks on completeness, correctness, and consistency",
        "Analysis Layer: Business metrics calculation for 15 required analysis questions",
        "Notification Layer: n8n workflow triggered via webhook with summary report",
        "Orchestration Layer: Apache Airflow DAG manages task dependencies and scheduling",
        "Archival Layer: Timestamped snapshots of raw and processed data",
    ]
    for s in steps:
        add_bullet(doc, s)
    add_para(doc, "Airflow handles extraction tasks in parallel, then runs merge, KNIME, validation, metrics, n8n notification, and archival tasks sequentially.")

    # 4. Extraction Process
    add_heading(doc, "4. Extraction Process")
    add_para(doc, "Each source has a dedicated Python extraction script:")
    for name, api, detail in [
        ("Arbeitnow", "arbeitnow.com/api/job-board-api", "Paginated requests up to 10 pages, 100 jobs per page. Handles JSON response with 'data' array and 'links' for pagination."),
        ("RemoteOK", "remoteok.com/api", "Single request returns all jobs as JSON array. First metadata entry is skipped. Salary fields (salary_min, salary_max) are included."),
        ("Himalayas", "himalayas.app/jobs/api/search", "Paginated requests with 'q=data' parameter. Returns 'jobs' array. Fields like companyName, applicationLink, skills are extracted."),
        ("RemoteJobs.org", "remotejobs.org/api/v1/jobs", "Multiple category queries (data-science, data-analytics, data-engineering, machine-learning, business-intelligence). Paginated up to 5 pages per category."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(detail)

    # 5. Schema Mapping
    add_heading(doc, "5. Schema Mapping")
    add_para(doc,
        "Each source returns different field names and structures. A mapping function per source "
        "standardizes all fields into 25 mandatory columns. Below is the mapping table:"
    )
    table = doc.add_table(rows=5, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Standard Field", "Arbeitnow", "RemoteOK", "Himalayas", "RemoteJobs.org"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    mappings = [
        ["title", "title", "position", "title", "title"],
        ["company_name", "company_name", "company", "companyName", "company.name"],
        ["location_raw", "location", "location", "locationRestrictions", "location"],
        ["remote_status", "remote (bool)", "Always Remote", "Always Remote", "Always Remote"],
    ]
    for i, row_data in enumerate(mappings):
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j].text = val

    # 6. KNIME Cleaning
    add_heading(doc, "6. KNIME Cleaning and Transformation")
    add_para(doc, "The KNIME workflow performs these operations in sequence:")
    knime_steps = [
        "Read merged CSV (CSV Reader node)",
        "Remove HTML tags and special characters from description, title, company_name, tags_raw (String Manipulation nodes)",
        "Deduplicate by source + job_url + title + company_name (Remove Duplicates node)",
        "Filter AI/ML/Data-related jobs using keyword matching on title and description (Rule Engine node)",
        "Classify jobs into categories: Data Analytics, Data Science, Data Engineering, AI/ML, BI, Analytics Engineering (Rule Engine node)",
        "Standardize remote_status into Remote/On-site/Hybrid/Unknown (Rule Engine node)",
        "Extract experience years from title and description, assign experience bracket (Rule Engine node)",
        "Keep only required columns (Column Filter node)",
        "Export cleaned dataset to CSV (CSV Writer node)",
    ]
    for s in knime_steps:
        add_bullet(doc, s)

    # 7. AI/ML/Data Filtering
    add_heading(doc, "7. AI/ML/Data Filtering Logic and Evidence")
    add_para(doc, "Filtering uses keyword matching against job title and description. The keywords span 6 categories:")
    categories_filter = [
        "Data Analytics: data analyst, data analytics, reporting analyst, product analyst, marketing analytics",
        "Data Science: data scientist, statistical modeling, predictive modeling, NLP, experimentation",
        "Data Engineering: data engineer, ETL, ELT, data pipeline, warehouse, lakehouse, dbt, Airflow",
        "AI/ML: machine learning, ML engineer, AI engineer, deep learning, LLM, computer vision, MLOps",
        "BI: BI analyst, business intelligence, Power BI, Tableau, dashboard developer",
        "Analytics Engineering: analytics engineer, dbt, semantic layer, metrics layer, data modeling",
    ]
    for c in categories_filter:
        add_bullet(doc, c)
    total_raw = metrics.get("total_raw", "N/A")
    total_clean = metrics.get("total_clean", "N/A")
    add_para(doc, f"Rows before filtering: {total_raw}")
    add_para(doc, f"Rows after filtering: {total_clean}")

    # 8. Salary Extraction
    add_heading(doc, "8. Salary Extraction and USD Conversion")
    add_para(doc,
        "Salary extraction handles multiple formats: structured salary_min/max fields, "
        "USD/EUR/GBP range patterns in descriptions, k-values ($85k), hourly rates ($40/hour -> $83,200/year), "
        "and monthly rates. All salaries are converted to annual USD using the Frankfurter exchange rate API "
        "(api.frankfurter.dev). Salary values of 0 are treated as missing."
    )

    # 9. Experience Brackets
    add_heading(doc, "9. Experience-Year Bracket Method")
    add_para(doc,
        "Experience requirements are extracted from job titles and descriptions using regex patterns "
        "(e.g., '3+ years', '1-3 years'). If no experience is mentioned, the bracket is 'Not mentioned'. "
        "Title-based inference is used: 'Junior' -> 0-1, 'Senior' -> 5-8+, 'Principal/Director' -> 8+."
    )

    # 10. Airflow DAG
    add_heading(doc, "10. Airflow DAG Explanation")
    add_para(doc,
        "The Airflow DAG (job_market_analytics_pipeline) orchestrates the complete pipeline with these tasks:"
    )
    dag_tasks = [
        "extract_arbeitnow, extract_remoteok, extract_himalayas, extract_remotejobs: Run in parallel",
        "merge_sources: Runs after all 4 extractions complete",
        "run_knime_workflow: Executes the KNIME cleaning workflow",
        "validate_clean_output: Checks row counts, columns, missing values, duplicates",
        "calculate_metrics: Computes all analysis metrics",
        "trigger_n8n_workflow: Sends metrics to n8n webhook",
        "archive_outputs: Saves timestamped copies",
    ]
    for t in dag_tasks:
        add_bullet(doc, t)
    add_para(doc, "The DAG runs weekly (Monday 6 AM) and has retry logic with 1 retry on failure.")

    # 11. n8n Workflow
    add_heading(doc, "11. n8n Workflow Explanation")
    add_para(doc,
        "The n8n workflow is triggered via a webhook endpoint (/webhook/job-market-alert). "
        "It receives the metrics payload from the Airflow DAG, formats it into a human-readable "
        "summary message, and responds back. The workflow includes: Webhook Trigger -> "
        "Format Message (JavaScript function) -> Respond to Webhook."
    )
    add_para(doc, "The notification includes: total cleaned jobs, jobs by source, remote/on-site ratio, "
        "count of 0-1 year jobs, average salary in USD, and pipeline run status.")

    # 12. Data Quality Checks
    add_heading(doc, "12. Data Quality Checks and Results")
    quality_items = [
        "API Response Check: Each source returns HTTP 200 and at least 1 job record",
        "Source Count Check: Raw job count per source recorded before merging",
        "Schema Check: Final dataset contains all 25 mandatory columns",
        "Duplicate Check: Duplicates counted and removed using source+url+title+company key",
        f"Relevance Filter: {total_raw} -> {total_clean} rows after AI/ML filtering",
        "Missing Value Check: Missing percentages calculated for title, company, url, description, dates, salary",
        "Date Check: Invalid dates fixed or marked; publication dates validated as YYYY-MM-DD",
        "Remote Status Check: All rows have Remote/On-site/Hybrid/Unknown",
        "Salary Check: Zero salaries converted to null; USD not calculated from missing",
        "Currency Conversion: Detected currencies logged with exchange rates and source API date",
        "Experience Bracket: Every row has a bracket or 'Not mentioned'",
        "Output File Check: clean_ai_ml_data_jobs.csv and metrics_summary.json exist and are non-empty",
    ]
    for q in quality_items:
        add_bullet(doc, q)

    # 13. Analysis Answers
    add_heading(doc, "13. Analysis Answers")

    raw_counts = metrics.get("total_raw_by_source", {})
    clean_counts = metrics.get("clean_by_source", {})
    total_raw = metrics.get("total_raw", 0)
    total_clean = metrics.get("total_clean", 0)

    qa_pairs = [
        ("Q1: Total jobs collected from each source before filtering",
         f"{', '.join(f'{k}: {v}' for k, v in raw_counts.items())}. Total: {total_raw}"),
        ("Q2: Jobs remaining after AI/ML/Data filtering",
         f"Total: {total_clean} (from {total_raw} raw, removed {total_raw - total_clean})"),
        ("Q3: Source with highest AI/ML/Data job count",
         f"{metrics.get('top_source', 'N/A')}"),
        ("Q4: Remote/on-site/hybrid/unknown ratio",
         f"{', '.join(f'{k}: {v}%' for k, v in metrics.get('remote_status_ratios', {}).items())}"),
        ("Q5: Remote/on-site ratio by source (limitation: RemoteOK, Himalayas, RemoteJobs.org are remote-only)",
         f"See metrics detail: {json.dumps(metrics.get('remote_by_source', {}), indent=2)}"),
        ("Q6: Jobs available for 0-1 year experience",
         f"{metrics.get('zero_one_year_jobs', 'N/A')} jobs"),
        ("Q7: Experience bracket distribution",
         f"{', '.join(f'{k}: {v}' for k, v in metrics.get('experience_brackets', {}).items())}"),
        ("Q8: Average salary in USD overall",
         f"${metrics.get('avg_salary_usd_overall', 'N/A')} (from {metrics.get('salary_count', 0)} records with salary data)"),
        ("Q9: Average salary by job category",
         f"{', '.join(f'{k}: ${v}' for k, v in metrics.get('avg_salary_by_category', {}).items())}"),
        ("Q10: Average salary by experience bracket",
         f"{', '.join(f'{k}: ${v}' for k, v in metrics.get('avg_salary_by_bracket', {}).items())}"),
        ("Q11: Most frequent skills",
         f"{', '.join(f'{s[0]} ({s[1]})' for s in metrics.get('top_skills', [])[:10])}"),
        ("Q12: Companies with most AI/ML/Data jobs",
         f"{', '.join(f'{c[0]} ({c[1]})' for c in metrics.get('top_companies', [])[:10])}"),
        ("Q13: Job category with highest openings",
         f"{max(metrics.get('job_category_counts', {}).items(), key=lambda x: x[1]) if metrics.get('job_category_counts') else 'N/A'}"),
        ("Q14: Source with highest salary coverage",
         f"{max(metrics.get('salary_coverage_by_source', {}).items(), key=lambda x: x[1]) if metrics.get('salary_coverage_by_source') else 'N/A'}"),
        ("Q15: Major data quality issues in raw sources",
         '; '.join(metrics.get('data_quality_issues', ['No issues reported']))),
    ]
    for q, a in qa_pairs:
        p = doc.add_paragraph()
        run_q = p.add_run(f"{q}: ")
        run_q.bold = True
        p.add_run(str(a))

    # 14. Challenges
    add_heading(doc, "14. Challenges and Solutions")
    challenges = [
        ("Schema Heterogeneity", "Each source has different field names and structures. Solution: per-source mapping functions with a unified standard schema of 25 columns."),
        ("Salary Format Variability", "Salaries appear as ranges, hourly rates, annual figures, in different currencies. Solution: regex patterns for 10+ formats, Frankfurter API for FX conversion, period multipliers for hourly/monthly."),
        ("Experience Inference", "Many jobs don't explicitly state required experience. Solution: keyword-based inference from titles (junior, senior, principal) combined with regex pattern matching in descriptions."),
        ("Rate Limiting", "Some APIs have request limits. Solution: conservative pagination (max 10 pages for Arbeitnow, 5 for others) with error handling and retries."),
        ("HTML in Descriptions", "Job descriptions contain HTML tags. Solution: regex-based HTML stripping in both Python and KNIME."),
    ]
    for title, desc in challenges:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # 15. Conclusion
    add_heading(doc, "15. Conclusion")
    add_para(doc,
        f"This pipeline successfully demonstrates an end-to-end data engineering workflow. "
        f"It collects job data from 4 sources, standardizes schemas, cleans and filters data in KNIME, "
        f"orchestrates the process with Airflow, notifies via n8n, and produces business insights. "
        f"The final dataset of {total_clean} AI/ML/Data jobs enables analysis of hiring demand, "
        f"salary trends, skill requirements, experience expectations, and remote work patterns "
        f"across the job market."
    )

    # Save
    report_path = BASE_DIR / "report" / "final_report.docx"
    doc.save(str(report_path))
    print(f"[REPORT] Saved to {report_path}")

if __name__ == "__main__":
    generate_report()
